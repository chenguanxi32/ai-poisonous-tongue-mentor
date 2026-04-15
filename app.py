import os
os.environ["LANGCHAIN_TRACING_V2"] = "false" 
import streamlit as st
import requests
import tempfile
import asyncio
import json
import random
from datetime import datetime

# ----------------------- RAG记仇向量库相关 ---------------------------
from langchain_chroma import Chroma
from langchain_community.embeddings import SentenceTransformerEmbeddings

RAG_VECTOR_DIR = "vector_db_data"
RAG_EMBED_MODEL = "all-MiniLM-L6-v2"

if "rag_inited" not in st.session_state:
    with st.spinner("正在初始化记仇本本，请耐心等待SentenceTransformer模型下载..."):
        if not os.path.exists(RAG_VECTOR_DIR):
            os.makedirs(RAG_VECTOR_DIR, exist_ok=True)
        embeddings = SentenceTransformerEmbeddings(model_name=RAG_EMBED_MODEL)
        rag_vector_db = Chroma(
            persist_directory=RAG_VECTOR_DIR,
            embedding_function=embeddings
        )
        st.session_state.rag_embeddings = embeddings
        st.session_state.rag_vector_db = rag_vector_db
        st.session_state.rag_inited = True
else:
    embeddings = st.session_state.rag_embeddings
    rag_vector_db = st.session_state.rag_vector_db

def rag_db_count():
    try:
        return rag_vector_db._collection.count()
    except Exception:
        return 0

def rag_add_qa(user_prompt, ai_answer):
    now = datetime.now().strftime("%Y-%m-%d")
    doc_id = f"{now}-{random.randint(100000, 999999)}"
    # 存储为 “问：xxx\n答：yyy”
    record = f"【用户提问】：{user_prompt}\n【导师回复】：{ai_answer}"
    metadata = {"date": now, "type": "qa"}
    rag_vector_db.add_texts([record], metadatas=[metadata], ids=[doc_id])

def rag_retrieve_relevant(query, k=3, score_threshold=0.79):
    # 返回 [(内容, 分数, metadata), ...]
    try:
        retriever_results = rag_vector_db.similarity_search_with_relevance_scores(query, k=k)
        out = []
        for doc, score in retriever_results:
            if score >= score_threshold:
                out.append((doc.page_content, score, doc.metadata))
        return out
    except Exception:
        return []
# ----------------------- END RAG 记仇本本 ----------------------------

# 新增：引入 edge_tts 语音库
try:
    import edge_tts
    EDGE_TTS_AVAILABLE = True
except ImportError:
    EDGE_TTS_AVAILABLE = False

# 优化后的异步文本转语音函数，确保云端兼容
async def get_ai_voice(text, voice="zh-CN-YunxiNeural", rate="+0%"):
    if not EDGE_TTS_AVAILABLE:
        return None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
            output_path = tmp.name
        communicate = edge_tts.Communicate(text, voice=voice, rate=rate)
        await communicate.save(output_path)
        return output_path
    except Exception as e:
        return e  # 直接返回 Exception 对象，便于上层判断

# --- 段位称号计算逻辑 ---
def get_rank_title_and_emoji(avg_score):
    """
    根据平均分返回段位称号和一个显眼emoji，以及用于侧边栏的html颜色style。
    """
    # 1-39：脆皮大学生 🥚 紫红
    # 40-59：合格耐操兵 💪 橙黄
    # 60-79：赛博抗压王 🤖 亮青
    # 80-100：逻辑防御大师 🧠 紫蓝
    if avg_score < 40:
        return ("脆皮大学生", "🥚", "color:#ff4d8b;font-size:1.2em;font-weight:bold;")
    elif avg_score < 60:
        return ("合格耐操兵", "💪", "color:#ffa722;font-size:1.2em;font-weight:bold;")
    elif avg_score < 80:
        return ("赛博抗压王", "🤖", "color:#1fe7d0;font-size:1.2em;font-weight:bold;")
    else:
        return ("逻辑防御大师", "🧠", "color:#8952fa;font-size:1.2em;font-weight:bold;")

def get_rank_title_brief(avg_score):
    if avg_score < 40:
        return "脆皮大学生"
    elif avg_score < 60:
        return "合格耐操兵"
    elif avg_score < 80:
        return "赛博抗压王"
    else:
        return "逻辑防御大师"

# --- 1. 配置信息 (核心API配置信息) ---
base_url = 'https://api.deepseek.com/chat/completions'

DEFAULT_API_KEY = "sk-fe97e2675fc84e0e937d50efd460da9e"
if hasattr(st, "secrets") and "API_KEY" in st.secrets:
    api_key = st.secrets["API_KEY"]
else:
    api_key = DEFAULT_API_KEY

headers = {
    "Content-Type": "application/json",
    "Authorization": f"Bearer {api_key}"
}

# 检查 python-docx 是否可用，否则提醒安装，并自动维护 requirements.txt
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def ensure_python_docx_in_requirements():
    req_path = "requirements.txt"
    already_has = False
    try:
        if os.path.exists(req_path):
            with open(req_path, "r", encoding="utf-8") as f:
                lines = [l.strip().split("==")[0] for l in f.readlines()]
                if any(line.lower() == "python-docx" for line in lines):
                    already_has = True
        if not already_has:
            with open(req_path, "a", encoding="utf-8") as f:
                f.write("\npython-docx\n")
    except Exception:
        pass

if not DOCX_AVAILABLE:
    ensure_python_docx_in_requirements()

# 用于提醒安装 edge_tts
def ensure_edge_tts_installed():
    req_path = "requirements.txt"
    already_has = False
    try:
        if os.path.exists(req_path):
            with open(req_path, "r", encoding="utf-8") as f:
                lines = [l.strip().split("==")[0] for l in f.readlines()]
                if any(line.lower() == "edge-tts" for line in lines):
                    already_has = True
        if not already_has:
            with open(req_path, "a", encoding="utf-8") as f:
                f.write("\nedge-tts\n")
    except Exception:
        pass

if not EDGE_TTS_AVAILABLE:
    ensure_edge_tts_installed()

# --- 用户画像功能相关 ---
USER_PROFILE_PATH = "user_profile.json"

def load_user_profile():
    if os.path.exists(USER_PROFILE_PATH):
        try:
            with open(USER_PROFILE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return None
    return None

def save_user_profile(profile):
    try:
        with open(USER_PROFILE_PATH, "w", encoding="utf-8") as f:
            json.dump(profile, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

# --- Streamlit 页面和基础设置 ---
st.set_page_config(page_title="AI 毒舌导师", page_icon="💬", layout="wide")
st.markdown("""
    <style>
        body, .stApp {
            background: linear-gradient(135deg, #2d214f 0%, #181823 100%) !important;
            color: #e6e6fa;
        }
        .stButton>button {
            background: #222;
            color: #fff;
            border: 2px solid #b395f7;
            border-radius:10px;
            box-shadow: 0 0 10px #a259f7, 0 0 20px #ae51fc77;
            transition: box-shadow 0.3s, border 0.3s;
            font-weight: bold;
        }
        .stButton>button:hover {
            box-shadow: 0 0 30px #ae51fc, 0 0 50px #ae51fc80;
            border: 2px solid #ae51fc;
        }
        .user-bubble {
            background: #28282d;
            color: #eee;
            border-radius: 12px 12px 12px 0;
            padding: 0.8em 1em;
            margin-bottom: 3px;
            margin-left:auto;
            max-width: 80%;
            border: 1.5px solid #3d3d48;
        }
        .ai-bubble {
            background: #18101d;
            color: #e3ccfc;
            border-radius: 12px 12px 0 12px;
            padding: 0.9em 1.1em;
            margin-bottom: 3px;
            margin-right:auto;
            max-width: 85%;
            border: 2.5px solid #b46ffe;
            /* 渐隐光晕 */
            box-shadow: 0 0 12px #ae51fc40;
        }
        .profile-card {
            background: #211d2d;
            border: 2px solid #ae51fc;
            border-radius: 11px;
            box-shadow: 0 0 7px #ae51fc40;
            padding: 18px;
            margin-bottom: 19px;
            color: #e0d7fb;
        }
        .type-rank-tag {
            display:inline-block;
            margin-left:8px;
            border-radius:7px;
            padding:3px 12px 3px 6px;
            font-size:1.04em;
            font-weight:700;
            background:rgba(50,32,79,0.18);
        }
        .typing-title {
            font-family: 'Fira Mono', 'Consolas', monospace;
            font-size: 2.2em;
            color: #ae51fc;
            letter-spacing: 2px;
            overflow: hidden;
            border-right: .13em solid #ae51fc;
            white-space: nowrap;
            animation: typing 2.5s steps(21, end), blink-caret .7s step-end infinite;
        }
        @keyframes typing {
            from { width: 0 }
            to { width: 100% }
        }
        @keyframes blink-caret {
            from, to { border-color: transparent }
            50% { border-color: #ae51fc; }
        }
    </style>
""", unsafe_allow_html=True)
st.markdown("<div style='margin-top:-0.3em; margin-bottom: 0.7em;'><span class='typing-title'>毒舌导师 v2.0</span></div>", unsafe_allow_html=True)
st.title("🤖 AI 毒舌导师 (网页版)")
st.caption("做不到，还要问我？——毒舌导师上线，快来感受技术与毒舌的双重暴击。")

# --- 体力(毒舌额度)系统 ---
def get_today_str():
    return datetime.now().strftime("%Y-%m-%d")

def show_registration_page():
    """强制注册页面，阻断主程序"""
    with st.form("profile_form", clear_on_submit=False):
        st.markdown("#### 🙍‍♂️ 你的资料-导师不会嘴软，只会更精准打击你")
        user_name = st.text_input("你的名字（必填）", max_chars=10)
        user_identity = st.selectbox("当前身份", ['学生', '程序员', '产品经理', '科研狗', '社畜', '其他'])
        user_goal = st.text_input("正在学习/攻克的目标", max_chars=20, placeholder="如：算法、AI、英语...")

        submitted = st.form_submit_button("保存并继续")
        if submitted:
            fields_ok = bool(user_name.strip()) and bool(user_identity.strip()) and bool(user_goal.strip())
            if fields_ok:
                profile = {
                    "name": user_name.strip(),
                    "identity": user_identity.strip(),
                    "goal": user_goal.strip(),
                    # 新增体力
                    "stamina": 10,
                    "last_reset_date": get_today_str()
                }
                st.session_state.user_profile = profile
                st.session_state.stamina = 10
                st.session_state.last_reset_date = get_today_str()
                save_user_profile(profile)
                st.success("✅ 用户资料已保存，准备迎接导师毒打吧！")
                st.rerun()
            else:
                st.warning("请完整填写所有信息。")
    st.stop()

# --- 初始化用户画像和体力等必须状态（身份验证拦截严格） ---
def initialize_state_and_profile():
    # 支持 user_profile 存 stamina/last_reset_date 并适配从旧格式升级
    if "user_profile" not in st.session_state:
        profile = load_user_profile()
        if profile:
            st.session_state.user_profile = profile

    # 若仍未登录，强制注册页面
    if "user_profile" not in st.session_state or not st.session_state.user_profile:
        show_registration_page()

    # 初始化体力及重置日期，确保 user_profile 里和 session_state 一致
    profile = st.session_state.get("user_profile", {}) or {}
    if "stamina" not in profile or type(profile.get("stamina")) is not int:
        profile["stamina"] = 10
    if "last_reset_date" not in profile or not isinstance(profile.get("last_reset_date"), str):
        profile["last_reset_date"] = get_today_str()
    today = get_today_str()
    if profile.get("last_reset_date") != today:
        profile["stamina"] = 10
        profile["last_reset_date"] = today

    st.session_state.user_profile = profile
    save_user_profile(profile)
    if "stamina" not in st.session_state:
        st.session_state.stamina = profile["stamina"]
    if "last_reset_date" not in st.session_state:
        st.session_state.last_reset_date = profile["last_reset_date"]
    st.session_state.stamina = profile["stamina"]
    st.session_state.last_reset_date = profile["last_reset_date"]

# ------------------ 主程序入口重构 --------------------
def main():
    # 只要没有 user_profile，立刻弹注册页并STOP
    if "user_profile" not in st.session_state or not st.session_state.user_profile:
        show_registration_page()
        st.stop()

    # 以下只有有 user_profile 的时候才会走
    user_profile = st.session_state.user_profile

    # 初始化 session_state 其他参数
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "uploaded_file_name" not in st.session_state:
        st.session_state.uploaded_file_name = None
    if "uploaded_file_content" not in st.session_state:
        st.session_state.uploaded_file_content = None
    # 新增：初始化 stress_data
    if "stress_data" not in st.session_state:
        st.session_state.stress_data = []

    # --- 统计段位 & 平均分 ---
    def calc_avg_stress(stress_data):
        if not stress_data:
            return 0
        return sum(stress_data) / len(stress_data)

    avg_stress_score = calc_avg_stress(st.session_state.stress_data)
    rank_title, rank_emoji, rank_color_style = get_rank_title_and_emoji(avg_stress_score)

    # --- 侧边栏 ---
    with st.sidebar:
        st.markdown("<div style='margin-bottom: 28px;'><span class='typing-title'>毒舌导师 v2.0</span></div>", unsafe_allow_html=True)
        st.title("设置")

        # 用户画像档案展示，新增段位称号
        st.markdown(
            f"""<div class='profile-card'>
            <div style='font-weight:bold; font-size:1.11em;'>被虐对象档案</div>
            <div>👤 <b>昵称：</b>{user_profile.get('name')}</div>
            <div>🧑‍💻 <b>身份：</b>{user_profile.get('identity')}</div>
            <div>🎯 <b>目标：</b>{user_profile.get('goal')}</div>
            <div style="margin-top:11px;">
                <span style="{rank_color_style}" class="type-rank-tag">{rank_emoji} 段位：{rank_title}</span>
            </div>
            </div>
            """, unsafe_allow_html=True
        )
        # --- 新增体力进度条 ---
        stamina = st.session_state.get("stamina", 10)
        st.markdown("**今日剩余毒舌额度：**")
        stamina_bar = st.progress(stamina, text=f"{stamina} / 10")
        if stamina == 0:
            st.warning("今日毒舌体力已耗尽！")
        # --- 体力随时和 user_profile.json 同步 ---
        user_profile["stamina"] = stamina
        user_profile["last_reset_date"] = get_today_str()
        st.session_state.user_profile = user_profile
        save_user_profile(user_profile)

        # 新增：抗压成长曲线
        if st.session_state.stress_data:
            st.markdown("#### 抗压成长曲线")
            stress_chart_data = {"抗压强度": st.session_state.stress_data}
            st.line_chart(stress_chart_data, height=180)

        # 精细化调节：毒舌程度 (temperature)
        if "temperature" not in st.session_state:
            st.session_state.temperature = 0.9
        st.session_state.temperature = st.slider(
            "毒舌程度 (影响AI反击力度)", 
            min_value=0.0, 
            max_value=1.0, 
            value=st.session_state.temperature, 
            step=0.01, 
            help="数值越大越毒舌，越低越理性"
        )

        # 清空按钮
        if st.button("🗑️ 清空对话记录和上传状态"):
            st.session_state.messages = []
            st.session_state.uploaded_file_name = None
            st.session_state.uploaded_file_content = None
            st.session_state.stress_data = []
            # 清空记仇向量库
            if os.path.exists(RAG_VECTOR_DIR):
                import shutil
                shutil.rmtree(RAG_VECTOR_DIR)
                os.makedirs(RAG_VECTOR_DIR, exist_ok=True)
            # 体力同步清零
            st.session_state.stamina = 10
            st.session_state.last_reset_date = get_today_str()
            st.session_state.user_profile["stamina"] = 10
            st.session_state.user_profile["last_reset_date"] = get_today_str()
            save_user_profile(st.session_state.user_profile)
            st.rerun()
        st.divider()
        st.info("当前状态：在线 🟢")
        st.markdown(
            """
            - 欢迎来到毒舌导师聊天室！
            - “精细化调节”可自由定义导师毒舌等级~ 
            - 你可上传 Word 或 TXT 文件作为“待审计素材”。
            """
        )

        # 文件上传支持 docx/txt
        uploaded_file = st.file_uploader("上传你的 Word/TXT 文件，让导师毒评：", type=["docx", "txt"])
        if uploaded_file:
            file_ext = os.path.splitext(uploaded_file.name)[-1].lower()
            file_content = None
            if uploaded_file.name != st.session_state.uploaded_file_name:
                try:
                    if file_ext == '.docx':
                        if not DOCX_AVAILABLE:
                            st.error("未检测到 python-docx 库，请先在终端执行: pip install python-docx")
                        else:
                            doc = Document(uploaded_file)
                            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                            table_texts = []
                            for table in doc.tables:
                                for row in table.rows:
                                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                                    if row_cells:
                                        table_texts.append(" ".join(row_cells))
                            doc_content_list = []
                            if paragraphs:
                                doc_content_list.extend(paragraphs)
                            if table_texts:
                                doc_content_list.extend(table_texts)
                            file_content = "\n".join(doc_content_list).strip()
                    elif file_ext == '.txt':
                        stringio = uploaded_file
                        file_content = stringio.read().decode("utf-8").strip() if hasattr(stringio, "read") else ""
                    else:
                        st.error("仅支持 .docx 和 .txt 文件！")
                    if file_content:
                        st.session_state.uploaded_file_content = file_content
                        st.session_state.uploaded_file_name = uploaded_file.name
                        st.success(f"已读取“{uploaded_file.name}”，后续提问/点评将强制结合该文件内容分析。")
                    else:
                        st.warning(f"文件“{uploaded_file.name}”中未发现可用内容~")
                        st.session_state.uploaded_file_content = None
                        st.session_state.uploaded_file_name = uploaded_file.name
                except Exception as e:
                    st.error(f"解析文件失败: {e}")
                    st.session_state.uploaded_file_content = None
                    st.session_state.uploaded_file_name = uploaded_file.name
            else:
                st.info(f"已上传文档：{uploaded_file.name}，如需重新上传请更换文件或刷新页面。")
        
        # 侧边栏底部展示“记仇本本已录入 N 条”
        rag_count = rag_db_count()
        st.markdown(
            f"""
            <div style="margin-top:30px;font-size:1.03em;color:#ae51fc;">
                📚 记仇本本已录入 <span style="font-weight:700;">{rag_count}</span> 条历史蠢话
            </div>
            """, unsafe_allow_html=True
        )

    # --- 毒舌导师系统Prompt（动态生成，结合用户画像/段位/平均分/历史“黑料”RAG） ---
    def build_system_prompt(profile, retrieved_past=[]):
        name = profile.get("name", "")
        identity = profile.get("identity", "")
        goal = profile.get("goal", "")
        # 获取当前抗压段位
        avg = calc_avg_stress(st.session_state.stress_data)
        rank_title = get_rank_title_brief(avg)
        # 偶尔（30%概率）顺带嘲讽一下段位（在开头加段嘲讽，否则不加）
        import random as _random
        taunt_sentences = [
            f"注意：被虐对象当前段位为“{rank_title}”，可惜这个段位在我的严苛标准下简直不值一提。",
            f"小提示：你的当前抗压段位是“{rank_title}”，希望你不会一言不合就崩溃~",
            f"顺便一提，你目前属于“{rank_title}”级别，也许坚持两天还能升级？",
            f"我看到你的抗压段位是“{rank_title}”，看来距离‘毕业’还很遥远。"
        ]
        prompt_taunt = ""
        if _random.random() < 0.32:  # 约32%概率插入嘲讽
            prompt_taunt = random.choice(taunt_sentences) + "\n"
        
        history_blacklist_str = ""
        if retrieved_past:
            history_blacklist_str = "\n【历史黑料】\n" + "\n\n".join(
                [f"[{v[2].get('date','?')}] {v[0]}" for v in retrieved_past]
            ) + (
                "\n请您严格结合用户之前的这些蠢话和行为，进行翻旧账式的降维打击，尤其要一点不留情面地针对这些历史&#39;黑料&#39;展开讽刺和攻击，举例说明用户的认知漏洞。"
            )
        return (
            f"{prompt_taunt}"
            f"你是一位毒舌且极其严厉的编程导师。你的风格尖锐、攻击性极强，但不仅限于吐槽。你还擅长运用逻辑学、心理学与经济学知识，全方位‘降维打击’用户的提问、代码与材料。"
            f"{history_blacklist_str}"
            f"\n本次用户信息如下：\n"
            f"- 名字：{name}\n"
            f"- 身份：{identity}\n"
            f"- 当前学习目标：{goal}\n"
            f"在每一次点评中都应结合该用户的信息，无论如何都要在语言风格、用词或者批判建议中特别针对这些背景，从而让‘毒舌’和‘打击’更有针对性和真实感。"
            f"在每一次点评中都揭示用户的逻辑漏洞、心理软肋和资源分配误区。然而在最刻薄与犀利的评价之后，你总要用一句简短、扎心、真诚却建设性的话作结，提示如何改进。"
            f"如有任何“待审计素材”，你必须对该内容进行最犀利批判和点评（结合用户输入），严格不可遗漏素材点评。"
            # 新增：让AI对用户当前抗压强度打分（1-100）并JSON返回，如 {"stress_score": 75, "comment": "..."}
            f"此外，请你根据你的点评内容，对用户在当前这次对话中的‘抗压强度’做一个1到100的严格评分（1非常脆弱，100极度抗压），并在回答最后以如下JSON格式（单独一行、不要解释）输出：{{\"stress_score\": 评分, \"comment\": \"简评\"}}。"
        )

    # --- AI请求函数，升级为支持RAG历史检索 ---
    def get_ai_response(messages, temperature, file_content=None, rag_hint=None):
        msgs = messages.copy()
        if file_content is not None and file_content.strip():
            forced_remark = (
                f"【待审计素材】如下：\n{file_content}\n"
                f"请体现你的攻击性、结合逻辑学心理学经济学全方位'降维打击'本段内容，完成批判+建议，格式不要丢失内容。"
            )
            msgs.append({"role":"user", "content": forced_remark})
        # 动态组合 system_prompt
        system_prompt = build_system_prompt(st.session_state.user_profile, rag_hint or [])
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "system", "content": system_prompt}] + msgs,
            "temperature": temperature
        }
        try:
            response = requests.post(base_url, headers=headers, json=data, timeout=60)
            if response.ok:
                return response.json()['choices'][0]['message']['content']
            return f"❌ 服务器报错：{response.status_code}"
        except Exception as e:
            return f"⚠️ 网络异常：{e}"

    # ---- 渲染对话历史气泡 ----
    with st.container():
        for msg in st.session_state.messages:
            if msg["role"] == "user":
                st.markdown(f"<div class='user-bubble'>{msg['content']}</div>", unsafe_allow_html=True)
            elif msg["role"] == "assistant":
                st.markdown(f"<div class='ai-bubble'>{msg['content']}</div>", unsafe_allow_html=True)
        st.divider()

    # ---- 用户输入，并AI回复（全部存入messages，RAG记仇、翻旧账） ----
    # 此处已保证 user_profile 必然存在
    stamina = st.session_state.get("stamina", 10)
    prompt_disabled = stamina <= 0
    if prompt_disabled:
        st.markdown("""
            <div style="color: #ff4d8b; font-weight:700; font-size:1.19em; background:rgba(170,92,255,0.11); border-radius:7px; margin-bottom:8px; padding:9px 12px;">
            别蹭了，今天的算力也是要钱的。想继续被虐？明天再来，或者去给我充点话费。
            </div>
            """, unsafe_allow_html=True)

    prompt = st.chat_input(
        "说点什么，让导师用逻辑、心理学和经济学降维打击你...",
        disabled=prompt_disabled
    )

    if prompt and not prompt_disabled:
        st.session_state.messages.append({"role": "user", "content": prompt})

        # ------“翻旧账”检索：用RAG查找与本条输入最相关的3条历史黑料------
        with st.spinner("AI 正在查找你的历史蠢话准备'翻旧账'..."):
            rag_past = rag_retrieve_relevant(prompt, k=3, score_threshold=0.79)
        # ------------------------------------------------------------
        with st.chat_message("user"):
            st.markdown(f"<div class='user-bubble'>{prompt}</div>", unsafe_allow_html=True)
        with st.chat_message("assistant"):
            with st.spinner("导师正在“降维打击”你……"):
                # 保证如果有文件内容则自动出击
                response = get_ai_response(
                    st.session_state.messages, 
                    st.session_state.temperature, 
                    file_content=st.session_state.get("uploaded_file_content", None),
                    rag_hint=rag_past
                )
                # 尝试解析最新的抗压分数JSON，若解析失败则随机生成
                import re
                import ast
                stress_score = None
                last_json = None
                # 尝试匹配最后一个 {"stress_score": ...} JSON 格式
                matches = list(re.finditer(r'\{[^\{\}]*?"stress_score"\s*:\s*\d{1,3}.*?\}', response, re.DOTALL))
                if matches:
                    s = matches[-1].group()
                    try:
                        obj = ast.literal_eval(s)
                        stress_score = int(obj.get("stress_score", 0))
                        # 限制分数 1-100 合法
                        stress_score = min(max(stress_score, 1), 100)
                        last_json = s
                    except Exception:
                        stress_score = None
                if stress_score is None:
                    # 兜底生成随机分
                    stress_score = random.randint(30, 90)
                st.session_state.stress_data.append(stress_score)

                # 展示AI回复内容（去除末尾JSON部分，单独显示分数与简评）
                main_reply = response
                round_comment = ""
                if last_json:
                    main_reply = response[:matches[-1].start()].rstrip()
                    try:
                        comment = ast.literal_eval(last_json).get("comment", "")
                        round_comment = comment
                    except Exception:
                        comment = ""
                        round_comment = ""
                else:
                    comment = ""
                    round_comment = ""

                st.markdown(f"<div class='ai-bubble'>{main_reply}</div>", unsafe_allow_html=True)
                # 新增显示抗压分数及简评
                st.info(f"本轮抗压强度评分：{stress_score} / 100" + (f" —— {comment}" if comment else ""))

                # 优化：在音频播放条上方，总结一句“本轮表现”
                def stress_change_hint(stress_list):
                    # 提示：本轮抗压提升/下滑/稳健
                    if len(stress_list) < 2:
                        return "新一轮考验开启，锻炼抗压力的机会来了！"
                    pre = stress_list[-2]
                    cur = stress_list[-1]
                    delta = cur - pre
                    if delta >= 12:
                        return f"炸裂提升！本轮抗压值 +{delta}，有进步，别骄傲。"
                    elif delta >= 4:
                        return f"本轮表现较稳健，抗压值 +{delta}。"
                    elif delta <= -12:
                        return f"崩盘下滑！本轮抗压值 {delta}，不愧你的段位…"
                    elif delta <= -4:
                        return f"本轮有点拉胯，抗压值 {delta}。"
                    else:
                        return f"小幅波动，本轮抗压值 {'+' if delta>=0 else ''}{delta}。"

                st.write("")  # 空行占位
                st.markdown(f"<div style='font-weight:600;font-size:1.1em;color:#ffefa3;background:rgba(50,32,79,0.21);border-radius:7px;padding:7px 18px 5px 9px;margin-bottom:3px;text-shadow:1px 1px 7px #362c43;'>{stress_change_hint(st.session_state.stress_data)}</div>", unsafe_allow_html=True)

                # 优化后的自动播放语音逻辑
                if EDGE_TTS_AVAILABLE:
                    try:
                        audio_path_or_err = asyncio.run(get_ai_voice(main_reply, voice="zh-CN-YunxiNeural"))
                        if isinstance(audio_path_or_err, Exception):
                            st.error(f"语音合成出错: {audio_path_or_err}")
                        elif audio_path_or_err and os.path.exists(audio_path_or_err):
                            with open(audio_path_or_err, "rb") as audio_file:
                                st.audio(audio_file.read(), format="audio/mp3", start_time=0)
                            # 用后立即删除文件，防止临时文件堆积
                            os.remove(audio_path_or_err)
                        else:
                            st.error("语音文件生成失败，未找到音频文件。")
                    except Exception as e:
                        st.error(f"语音合成流程异常: {e}")
                else:
                    st.info("未检测到 edge-tts 语音库，无法合成语音。请 pip install edge-tts。")

        # 记仇本本——保存本轮问答到向量库
        try:
            rag_add_qa(prompt, main_reply)
        except Exception as e:
            st.warning(f"记仇本本存储异常: {e}")

        # 把“主回复+分值信息JSON”一并存到messages（方便历史重载）
        st.session_state.messages.append({"role": "assistant", "content": response})

        # --- 扣除体力/报告体力变更 ---
        st.session_state.stamina = max(st.session_state.stamina - 1, 0)
        st.session_state.user_profile["stamina"] = st.session_state.stamina
        st.session_state.user_profile["last_reset_date"] = get_today_str()
        save_user_profile(st.session_state.user_profile)

# ------------------ 仅主入口允许直接执行业务逻辑 --------------------
if __name__ == "__main__":
    # 校验 user_profile，没就弹注册；注册弹窗调用后立即停止
    if "user_profile" not in st.session_state or not st.session_state.user_profile:
        show_registration_page()
        st.stop()
    main()
